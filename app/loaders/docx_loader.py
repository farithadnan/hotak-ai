"""Module to load docx documents."""

from typing import List
from pathlib import Path
from langchain_core.documents import Document
from ..config.settings import UPLOADS_DIRECTORY
from ..utils.logger import setup_logger

from docx import Document as DocxDocument

logger = setup_logger(__name__)

_ALLOWED_ROOT = UPLOADS_DIRECTORY.resolve()


def load_docx_document(file_path: str) -> List[Document]:
    """
    Load a DOCX document from the given file path.

    Args:
        file_path (str): The path to the DOCX file to load.

    Returns:
        List[Document]: Loaded documents.
    """
    # Construct path from trusted root using only the filename component.
    safe_path = (_ALLOWED_ROOT / Path(file_path).name).resolve()
    if not safe_path.is_relative_to(_ALLOWED_ROOT):
        logger.warning("Blocked path traversal attempt in DOCX loader: %s", file_path)
        raise ValueError("Access denied: file must be inside the uploads directory.")

    try:
        logger.info(f"Loading DOCX document from: {safe_path}")

        # Read the DOCX file
        docx = DocxDocument(str(safe_path))
        full_text = []
        for para in docx.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        content = "\n\n".join(full_text)

        # Create a single Document
        doc = Document(
            page_content=content,
            metadata={
                "source": str(safe_path),
                "file_name": safe_path.name,
                "source_type": "docx",
            }
        )

        para_count = len(full_text)
        char_count = len(content)
        logger.info(f"Loaded DOCX: {para_count} paragraphs, {char_count} characters")

        return [doc]
    except Exception as e:
        logger.error(f"Failed to load DOCX document from {safe_path}: {e}")
        raise